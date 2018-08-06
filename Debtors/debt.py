import pandas as pd
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
df = pd.read_excel("D:/R/должники/Дома.xlsx")

date="сегодняшняя дата"
mesyac="какой месяц"

def select_house(house):
    return pd.DataFrame(df.loc[df.Дом == house]).reset_index(drop=True)


def select_house_and_do_stuff(house):
    doc = docx.Document('D:/R/должники/Шаблон.docx')
    selected_house = select_house(house)
    house_name = selected_house.Дом[0].replace("/", "-")
    row_num = len(selected_house)
    doctable = doc.add_table(row_num, 3)
    doctable.style = 'Table Grid'
    for row, columns in selected_house.reset_index(drop=True).iterrows():
        doctable.rows[row].cells[0].text = columns["Дом"]
        doctable.rows[row].cells[1].text = columns["Лицевой счет.Помещение"]
        doctable.rows[row].cells[2].text = str(columns["Итого"])
    doc.add_paragraph('')
    doc.add_paragraph('          текст.')
    doc.add_paragraph('          текст')
    doc.add_paragraph('          текст')
    doc.add_paragraph('          текст')
    doc.add_paragraph('')
    doc.add_paragraph('текст')
    doc.add_paragraph('текст')
    for paragraph in doc.paragraphs:
        if paragraph.text == "          текст":
            paragraph.runs[0].font.size = Pt(12)
            paragraph.runs[0].font.name = 'Times New Roman'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif paragraph.text == "          текст":
            paragraph.runs[0].font.size = Pt(12)
            paragraph.runs[0].font.name = 'Times New Roman'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif paragraph.text == "          текст":
            paragraph.runs[0].font.size = Pt(12)
            paragraph.runs[0].font.name = 'Times New Roman'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif paragraph.text == "          текст":
            paragraph.runs[0].font.size = Pt(12)
            paragraph.runs[0].font.name = 'Times New Roman'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif paragraph.text == "текст":
            paragraph.runs[0].font.size = Pt(12)
            paragraph.runs[0].font.name = 'Times New Roman'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif paragraph.text == 'текст':
            paragraph.runs[0].font.size = Pt(12)
            paragraph.runs[0].font.name = 'Times New Roman'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for num, row in enumerate(doc.tables[0].rows):
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
    return doc.save("D:/R/должники/" + house_name + ".docx")

#select_house_and_do_stuff()

####################################################################################

def all_houses():
    houses = df.Дом.drop_duplicates()
    return houses

for house in all_houses(): # ЗАПУСКАТЬ ЭТО ДЛЯ МАССОВОЙ ОБРАБОТКИ
    select_house_and_do_stuff(house)

