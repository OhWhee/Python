from PyQt5 import QtCore, QtGui, QtWidgets

class Ui_Dialog(object):
    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.resize(400, 135)
        self.OpenBtn = QtWidgets.QPushButton(Dialog)
        self.OpenBtn.setGeometry(QtCore.QRect(10, 30, 75, 23))
        self.OpenBtn.setObjectName("OpenBtn")
        self.OpenBtn.clicked.connect(self.openButton)
        self.SaveBtn = QtWidgets.QPushButton(Dialog)
        self.SaveBtn.setGeometry(QtCore.QRect(10, 60, 75, 23))
        self.SaveBtn.setObjectName("SaveBtn")
        self.SaveBtn.clicked.connect(self.saveButton)
        self.StrOpen = QtWidgets.QLineEdit(Dialog)
        self.StrOpen.setGeometry(QtCore.QRect(100, 30, 281, 20))
        self.StrOpen.setObjectName("StrOpen")
        self.StrSave = QtWidgets.QLineEdit(Dialog)
        self.StrSave.setGeometry(QtCore.QRect(100, 60, 281, 20))
        self.StrSave.setObjectName("StrSave")
        self.ConvertTable = QtWidgets.QPushButton(Dialog)
        self.ConvertTable.setGeometry(QtCore.QRect(100, 90, 281, 23))
        self.ConvertTable.setObjectName("ConvertTable")
        self.ConvertTable.clicked.connect(self.makeCool)

        self.retranslateUi(Dialog)
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Dialog"))
        self.OpenBtn.setText(_translate("Dialog", "Открыть"))
        self.SaveBtn.setText(_translate("Dialog", "Сохранить"))
        self.ConvertTable.setText(_translate("Dialog", "Сделать хорошо"))
        
    def openButton(self):
        #self.StrOpen.setText(QtWidgets.QFileDialog.getExistingDirectory())
        filter = "Word (*.docx)"
        filename = QtWidgets.QFileDialog.getOpenFileName()
        self.StrOpen.setText(filename[0])
        
    def saveButton(self):
        self.StrSave.setText(QtWidgets.QFileDialog.getExistingDirectory())
        
    def makeCool(self):
        filepathSave = self.StrSave.text()
        filepathOpen = self.StrOpen.text()
        import docx
        from docx.shared import Pt
        doc = docx.Document(filepathOpen)
        tables = doc.tables
        
        tables[5].rows[0].cells[0].text = "текст"
        tables[5].rows[0].cells[3].text = "текст"
        tables[5].rows[1].cells[0].text = "текст"
        tables[5].rows[1].cells[1].text = "текст"
        tables[5].rows[2].cells[0].text = "текст"
        tables[5].rows[2].cells[1].text = "текст"
        tables[5].rows[3].cells[0].text = "текст"
        tables[5].rows[3].cells[1].text = "текст"
        tables[5].rows[4].cells[0].text = "текст"
        tables[5].rows[4].cells[1].text = "текст"
        tables[5].rows[5].cells[0].text = "текст"
        tables[5].rows[5].cells[1].text = "текст"
        tables[5].rows[6].cells[0].text = "текст"
        tables[5].rows[6].cells[1].text = "текст"
        tables[5].rows[7].cells[0].text = "текст"
        tables[5].rows[7].cells[1].text = "текст"
        tables[5].rows[8].cells[0].text = "текст"
        tables[5].rows[8].cells[1].text = "текст"
        tables[5].rows[9].cells[0].text = "текст"
        tables[5].rows[9].cells[1].text = "текст"
        tables[5].rows[10].cells[0].text = "текст"
        tables[5].rows[10].cells[2].text = "текст"
        
        tables[5].rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[0].cells[3].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[2].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[2].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[4].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[4].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[5].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[5].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[6].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[6].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[7].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[7].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[8].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[8].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[9].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[9].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[10].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[5].rows[10].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        
        tables[5].rows[0].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[0].cells[3].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[1].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[1].cells[1].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[2].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[2].cells[1].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[3].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[3].cells[1].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[4].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[4].cells[1].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[5].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[5].cells[1].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[6].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[6].cells[1].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[7].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[7].cells[1].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[8].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[8].cells[1].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[9].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[9].cells[1].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[10].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[5].rows[10].cells[2].paragraphs[0].runs[0].font.italic = True
        
        tables[5].rows[0].cells[3].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[1].cells[1].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[2].cells[1].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[3].cells[1].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[4].cells[1].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[5].cells[1].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[6].cells[1].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[7].cells[1].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[8].cells[1].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[9].cells[1].paragraphs[0].runs[0].font.bold = True
        tables[5].rows[10].cells[2].paragraphs[0].runs[0].font.bold = True
        
        tables[2].rows[12].cells[0].text = "текст"
        tables[2].rows[13].cells[9].text = "текст"
        tables[2].rows[14].cells[9].text = "текст"
        tables[2].rows[15].cells[9].text = "текст"
        tables[2].rows[16].cells[9].text = "текст"
        tables[2].rows[17].cells[9].text = "текст"
        tables[2].rows[18].cells[9].text = "текст"
        
        tables[2].rows[12].cells[0].paragraphs[0].runs[0].font.bold = True
        tables[2].rows[13].cells[9].paragraphs[0].runs[0].font.bold = True
        tables[2].rows[14].cells[9].paragraphs[0].runs[0].font.bold = True
        tables[2].rows[15].cells[9].paragraphs[0].runs[0].font.bold = True
        tables[2].rows[16].cells[9].paragraphs[0].runs[0].font.bold = True
        tables[2].rows[17].cells[9].paragraphs[0].runs[0].font.bold = True
        tables[2].rows[18].cells[9].paragraphs[0].runs[0].font.bold = True
        
        tables[2].rows[12].cells[0].paragraphs[0].runs[0].font.italic = True
        tables[2].rows[13].cells[9].paragraphs[0].runs[0].font.italic = True
        tables[2].rows[14].cells[9].paragraphs[0].runs[0].font.italic = True
        tables[2].rows[15].cells[9].paragraphs[0].runs[0].font.italic = True
        tables[2].rows[16].cells[9].paragraphs[0].runs[0].font.italic = True
        tables[2].rows[17].cells[9].paragraphs[0].runs[0].font.italic = True
        tables[2].rows[18].cells[9].paragraphs[0].runs[0].font.italic = True
        
        tables[2].rows[12].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        tables[2].rows[13].cells[9].paragraphs[0].runs[0].font.size = Pt(10)
        tables[2].rows[14].cells[9].paragraphs[0].runs[0].font.size = Pt(10)
        tables[2].rows[15].cells[9].paragraphs[0].runs[0].font.size = Pt(10)
        tables[2].rows[16].cells[9].paragraphs[0].runs[0].font.size = Pt(10)
        tables[2].rows[17].cells[9].paragraphs[0].runs[0].font.size = Pt(10)
        tables[2].rows[18].cells[9].paragraphs[0].runs[0].font.size = Pt(10)
        
        tables[3].rows[2].cells[1].text = "текст"
        tables[3].rows[2].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        tables[3].rows[2].cells[1].paragraphs[0].runs[0].font.bold = True
        
        for paragraph in doc.paragraphs:
            if  paragraph.text == "текст":
                paragraph.text = "текст"
                paragraph.runs[0].font.size = Pt(10)
                print(paragraph.text)
            elif paragraph.text == "текст":
                 paragraph.text = "текст"
                 paragraph.runs[0].font.size = Pt(12)
                 paragraph.runs[0].font.bold = True
                 print(paragraph.text)
            elif paragraph.text == "текст":
                 paragraph.text = "текст" 
                 paragraph.runs[0].font.size = Pt(12)
                 paragraph.runs[0].font.bold = True
                 print(paragraph.text) 
        
        filename = str(doc.paragraphs[0].text)
        filedate = str(tables[0].rows[0].cells[1].text)
        tblformat = ".docx"
        doc.save(filepathSave + '\\' + filename + ' ' + filedate + tblformat)


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    Dialog = QtWidgets.QDialog()
    ui = Ui_Dialog()
    ui.setupUi(Dialog)
    Dialog.show()
    sys.exit(app.exec_())
    
