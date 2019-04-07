# -*- coding: utf-8 -*-
import sqlite3
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import finmonxml

#debtor = ('8',)
#receiver = ('17',)
#payer = ('7',)
#
#participants = {'debtor' : '8', 'receiver' : '17', 'payer' : '7'}

class ParticipantObject():
    
    def __init__(self, org_name, type_id, inn, kpp, okpo, okved, ogrn, reg_organization, reg_date, country_id_ur, okato_subject_id_ur, area_ur, locality_ur,
                 street_ur, house_ur, corpus_ur, office_ur, country_id_post, okato_subject_id_post, area_post, locality_post, street_post, 
                 house_post, corpus_post, office_post, raschetniy_schet, bik, bank_name):
        self.org_name = org_name
        self.type_id = type_id
        self.inn = inn
        self.kpp = kpp
        self.okpo = okpo
        self.okved = okved
        self.ogrn = ogrn
        self.reg_organization = reg_organization
        self.reg_date = reg_date
        self.country_id_ur = country_id_ur
        self.okato_subject_id_ur = okato_subject_id_ur
        self.area_ur = area_ur
        self.locality_ur = locality_ur
        self.street_ur = street_ur
        self.house_ur = house_ur
        self.corpus_ur = corpus_ur
        self.office_ur = office_ur
        self.country_id_post = country_id_post
        self.okato_subject_id_post = okato_subject_id_post
        self.area_post = area_post
        self.locality_post = locality_post
        self.street_post = street_post
        self.house_post = house_post
        self.corpus_post = corpus_post
        self.office_post = office_post
        self.raschetniy_schet = raschetniy_schet
        self.bik = bik
        self.bank_name = bank_name
        
    def make_address_for_template_ul_ur(self):
        return '{} {} {} {} {}'.format(self.locality_ur, self.street_ur, self.house_ur, self.corpus_ur, self.office_ur)
    
    def make_address_for_template_ul_post(self):
        return '{} {} {} {} {}'.format(self.locality_post, self.street_post, self.house_post, self.corpus_post, self.office_post)


class OperationObjectFactory():
        
        def create_last_operation_object(self):
            conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
            cur = conn.cursor()
            annual_number = ''
            cur.execute("""SELECT * FROM operations WHERE ROWID IN (SELECT max(ROWID) FROM operations)""")
            rows = cur.fetchall()
            operation_date = ''
            detection_date = ''
            daily_number = ''
            annual_number = ''
            operation_number = ''
            amount = ''
            payment_purpose = ''
            payment_class = ''
            payment_number = ''
            debtor = ''
            payer = ''
            receiver = ''
            contract = ''
            operation_list = [operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract]
            i=0
            for r in rows[0][1:]:
                operation_list[0+i] = r
                i+=1
            operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract = [str(e) for e in operation_list]
            return OperationObject(operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract)  

        def create_operation_object_by_id(self, operation_id):
            conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
            cur = conn.cursor()
            annual_number = ''
            cur.execute("""SELECT * FROM operations WHERE operation_id = ?""", (operation_id,))
            rows = cur.fetchall()
            operation_date = ''
            detection_date = ''
            daily_number = ''
            annual_number = ''
            operation_number = ''
            amount = ''
            payment_purpose = ''
            payment_class = ''
            payment_number = ''
            debtor = ''
            payer = ''
            receiver = ''
            contract = ''
            operation_list = [operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract]
            i=0
            for r in rows[0][1:]:
                operation_list[0+i] = r
                i+=1
            operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract = [str(e) for e in operation_list]
            return OperationObject(operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract)


        


class OperationObject():
    
    def __init__(self, operation_date=None, detection_date=None, daily_number=None, annual_number=None,
                 operation_number=None, amount=None, payment_purpose=None, payment_class=None, payment_number=None, payer=None,
                 debtor=None, receiver=None, contract=None):
        self.operation_date = operation_date
        self.detection_date = detection_date
        self.daily_number = daily_number
        self.annual_number = annual_number
        self.operation_number = operation_number
        self.amount = amount
        self.payment_purpose = payment_purpose
        self.payment_class = payment_class
        self.payment_number = payment_number
        self.payer = payer
        self.debtor = debtor
        self.receiver = receiver
        self.contract = contract
        
        
    def get_requisites_for_ul(self, participant):
        conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
        cur = conn.cursor()
        cur.execute("""SELECT 
                            organizations.org_name,
                            organizations.type_id,
                            requisites_ul.inn,
                            requisites_ul.kpp,
                            requisites_ul.okpo,
                            requisites_ul.okved,
                            requisites_ul.ogrn,
                            requisites_ul.reg_organization,
                            requisites_ul.reg_date,
                            requisites_ul.country_id_ur,
                            requisites_ul.okato_subject_id_ur,
                            requisites_ul.area_ur,
                            requisites_ul.locality_ur,
                            requisites_ul.street_ur,
                            requisites_ul.house_ur,
                            requisites_ul.corpus_ur,
                            requisites_ul.office_ur,
                            requisites_ul.country_id_post,
                            requisites_ul.okato_subject_id_post,
                            requisites_ul.area_post,
                            requisites_ul.locality_post,
                            requisites_ul.street_post,
                            requisites_ul.house_post,
                            requisites_ul.corpus_post,
                            requisites_ul.office_post,
                            accounts.raschetniy_schet,
                            accounts.bik,
                            banks.bank_name
                            FROM organizations
                            JOIN requisites_ul ON organizations.org_id = requisites_ul.org_id
                            JOIN accounts ON organizations.org_id = accounts.org_id
                            JOIN banks ON banks.bik = accounts.bik
                            WHERE organizations.org_id = ?""", (participant,))
        rows = cur.fetchall()
        obj = finmonxml.Requisites(*rows[0])
        return obj      
    
    def get_requisites_by_type_id(self, participant):
        conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
        cur = conn.cursor()
        cur.execute("""SELECT 
                            organizations.type_id
                            FROM organizations
                            WHERE organizations.org_id = ?""", (participant,))
        rows = cur.fetchall()
        conn.close()
        return self.get_full_requisites(participant=participant, type_id=rows[0][0])
    
    def get_full_requisites(self, participant, type_id=None):
        conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
        cur = conn.cursor()
        obj = ''
        if type_id == 0:
            cur.execute("""SELECT 
                                organizations.org_name,
                                organizations.type_id,
                                requisites_ip.first_name,
                                requisites_ip.middle_name,
                                requisites_ip.last_name,
                                requisites_ip.inn,
                                requisites_ip.country_citizenship_id,
                                requisites_ip.birth_date,
                                requisites_ip.birth_place_country_id,
                                requisites_ip.birth_place_okato_subject_id,
                                requisites_ip.birth_place_area,
                                requisites_ip.birth_place_locality,
                                requisites_ip.document_type,
                                requisites_ip.document_series,
                                requisites_ip.document_number,
                                requisites_ip.document_date,
                                requisites_ip.document_given_by,
                                requisites_ip.department_number,
                                requisites_ip.country_id_registration,
                                requisites_ip.okato_subject_id_registration,
                                requisites_ip.area_registration,
                                requisites_ip.locality_registration,
                                requisites_ip.street_registration,
                                requisites_ip.house_registration,
                                requisites_ip.corpus_registration,
    							requisites_ip.office_registration,
    							requisites_ip.country_id_residence,
    							requisites_ip.okato_subject_id_residence,
    							requisites_ip.area_residence,
    							requisites_ip.locality_residence,
    							requisites_ip.street_residence,
    							requisites_ip.house_residence,
    							requisites_ip.corpus_residence,
    							requisites_ip.office_registration,
                                accounts.raschetniy_schet,
                                accounts.bik,
                                banks.bank_name
                                FROM organizations
                                JOIN requisites_ip ON organizations.org_id = requisites_ip.org_id
                                JOIN accounts ON organizations.org_id = accounts.org_id
                                JOIN banks ON banks.bik = accounts.bik
                                WHERE organizations.org_id = ?""", (participant,))
            rows = cur.fetchall()
            return finmonxml.ParticipantIP(*rows[0])
        elif type_id == 1:
            cur.execute("""SELECT 
                                organizations.org_name,
                                organizations.type_id,
								requisites_ul.inn,
								requisites_ul.kpp,
								requisites_ul.okpo,
								requisites_ul.okved, 
								requisites_ul.ogrn, 
								requisites_ul.reg_organization, 
								requisites_ul.reg_date, 
								requisites_ul.country_id_ur, 
								requisites_ul.okato_subject_id_ur, 
								requisites_ul.area_ur, 
								requisites_ul.locality_ur, 
								requisites_ul.street_ur, house_ur,
								requisites_ul.corpus_ur,
								requisites_ul.office_ur,
								requisites_ul.country_id_post, 
								requisites_ul.okato_subject_id_post, 
								requisites_ul.area_post, 
								requisites_ul.locality_post, 
								requisites_ul.street_post, 
								requisites_ul.house_post, 
								requisites_ul.corpus_post, 
								requisites_ul.office_post,
                                accounts.raschetniy_schet,
                                accounts.bik,
                                banks.bank_name
                                FROM organizations
                                JOIN requisites_ul ON organizations.org_id = requisites_ul.org_id
                                JOIN accounts ON organizations.org_id = accounts.org_id
                                JOIN banks ON banks.bik = accounts.bik
                                WHERE organizations.org_id = ?""",(participant,))
            rows = cur.fetchall()
            return finmonxml.ParticipantUL(*rows[0])
        elif type_id == 2:
            cur.execute("""SELECT 
                                organizations.org_name,
                                organizations.type_id,
                                requisites_fl.first_name,
                                requisites_fl.middle_name,
                                requisites_fl.last_name,
                                requisites_fl.inn,
                                requisites_fl.country_citizenship_id,
                                requisites_fl.birth_date,
                                requisites_fl.birth_place_country_id,
                                requisites_fl.birth_place_okato_subject_id,
                                requisites_fl.birth_place_area,
                                requisites_fl.birth_place_locality,
                                requisites_fl.document_type,
                                requisites_fl.document_series,
                                requisites_fl.document_number,
                                requisites_fl.document_date,
                                requisites_fl.document_given_by,
                                requisites_fl.department_number,
                                requisites_fl.country_id_registration,
                                requisites_fl.okato_subject_id_registration,
                                requisites_fl.area_registration,
                                requisites_fl.locality_registration,
                                requisites_fl.street_registration,
                                requisites_fl.house_registration,
                                requisites_fl.corpus_registration,
    							requisites_fl.office_registration,
    							requisites_fl.country_id_residence,
    							requisites_fl.okato_subject_id_residence,
    							requisites_fl.area_residence,
    							requisites_fl.locality_residence,
    							requisites_fl.street_residence,
    							requisites_fl.house_residence,
    							requisites_fl.corpus_residence,
    							requisites_fl.office_registration,
                                accounts.raschetniy_schet,
                                accounts.bik,
                                banks.bank_name
                                FROM organizations
                                JOIN requisites_fl ON organizations.org_id = requisites_fl.org_id
                                JOIN accounts ON organizations.org_id = accounts.org_id
                                JOIN banks ON banks.bik = accounts.bik
                                WHERE organizations.org_id = ?""", (participant,))
            rows = cur.fetchall()
            return finmonxml.ParticipantFL(*rows[0])    
        
        
    def make_next_annual_number(self):
        print('Предыдущий номер ФЭС - {}, следующий - {}'.format(self.annual_number, str(int(self.annual_number)+1)))
        return str(int(self.annual_number)+1)


    def make_next_operation_number(self):
        annual_number = int(self.make_next_annual_number())
        now = datetime.datetime.now()
        formatter = '{:08d}'.format
        operation_number = '{}_3444213503_344401001_{}'.format(now.year, formatter(annual_number))
        return operation_number

    def make_filename(self, date):
        daily_number = int(self.daily_number)
#        now = datetime.datetime.now()
        formatter = '{:08d}'.format
        filename = 'FM01_2_3444213503344401001_{}_01{}'.format(date, formatter(daily_number))
        return filename

    def get_last_detection_date(self):
        return self.detection_date
    
    
    def make_new_daily_number(self):
        now = datetime.datetime.now()
        detection_date = self.get_last_detection_date()
        daily_number = int(self.get_last_daily_number())
        if str(now.strftime("%d.%m.%Y")) != detection_date:
            daily_number = 1
        else:
            daily_number += 1
        return daily_number


    def get_last_daily_number(self):
        return self.daily_number
    
    
    def make_template_ul(self, operation_id, date=datetime.datetime.now().strftime("%Y%m%d")):
            with open('D:/python/FinMon/template_ul.docx', 'rb') as file:
                document = Document(file)
                paragraphs = document.paragraphs
                tables = document.tables
                last_operation = self.create_operation_object_by_id(operation_id)
                payer_id = last_operation.payer
                debtor_id = last_operation.debtor
                receiver_id = last_operation.receiver
                
                payer = last_operation.ge(participant=payer_id)
                debtor = last_operation.get_requisites_for_ul(participant=debtor_id)
                receiver = last_operation.get_requisites_for_ul(participant=receiver_id)
                
                paragraphs[1].text = 'ВНУТРЕННЕЕ СООБЩЕНИЕ (УВЕДОМЛЕНИЕ)\nо совершённой операции, подлежащей обязательному контролю, либо необычной (подозрительной) сделке № {} / {}'.format('{:08d}'.format(int(last_operation.daily_number)), '{:08d}'.format(int(last_operation.annual_number)))
                paragraphs[1].runs[0].font.size = Pt(14)
                paragraphs[1].runs[0].font.bold = True
                paragraphs[1].runs[0].font.name = 'Times New Roman'
                
                paragraphs[2].text = """Мною, Максименко Егором Вячеславовичем, ведущим инженером ООО «РЦ ЮО», в ходе осуществления своей профессиональной деятельности {}. выявлена операция, совершенная {}, связанная с поступлением денежных средств на р/с{} {} от {} на основании платежного документа №{} на сумму {} (плательщик {}, основание платежа: {})""".format(last_operation.detection_date, last_operation.operation_date, receiver.raschetniy_schet, receiver.org_name, payer.org_name, last_operation.payment_number, last_operation.amount, payer.org_name, last_operation.payment_purpose)
                paragraphs[2].runs[0].font.name = 'Times New Roman'
                paragraphs[2].runs[0].font.size = Pt(12)
                
                paragraphs[5].text = '{}'.format(receiver.org_name)
                paragraphs[5].runs[0].font.size = Pt(12)
                paragraphs[5].runs[0].font.bold = True
                paragraphs[5].runs[0].font.name = 'Times New Roman'
                paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                paragraphs[6].text = '{}'.format(payer.org_name)
                paragraphs[6].runs[0].font.size = Pt(12)
                paragraphs[6].runs[0].font.bold = True
                paragraphs[6].runs[0].font.name = 'Times New Roman'
                paragraphs[6].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                paragraphs[7].text = '{}'.format(debtor.org_name)
                paragraphs[7].runs[0].font.size = Pt(12)
                paragraphs[7].runs[0].font.bold = True
                paragraphs[7].runs[0].font.name = 'Times New Roman'
                paragraphs[7].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                tables[1].rows[0].cells[0].paragraphs[0].text = '{}г.'.format(last_operation.detection_date)
                tables[1].rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                tables[1].rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                
                tables[2].rows[0].cells[0].paragraphs[0].text = """{}\nИНН: {}\nКПП: {}\nОГРН: {}\nОКПО: {}\nОКВЭД: {}\nР/c: {}\nБИК: {}\nНаименование банка: {}\nНаименование регистрирующего органа: {}\nДата регистрации: {}\nЮридический адрес: {}\nФактический адрес: {}""".format(payer.org_name, payer.inn, payer.kpp, payer.ogrn, payer.okpo, payer.okved, payer.raschetniy_schet, payer.bik, payer.bank_name, payer.reg_organization, payer.reg_date, payer.make_address_for_template_ul_ur(), payer.make_address_for_template_ul_post())
                tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
                tables[2].rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                tables[2].rows[0].cells[1].paragraphs[0].text = """{}\nИНН: {}\nКПП: {}\nОГРН: {}\nОКПО: {}\nОКВЭД: {}\nР/c: {}\nБИК: {}\nНаименование банка: {}\nНаименование регистрирующего органа: {}\nДата регистрации: {}\nЮридический адрес: {}\nФактический адрес: {}""".format(debtor.org_name, debtor.inn, debtor.kpp, debtor.ogrn, debtor.okpo, debtor.okved, debtor.raschetniy_schet, debtor.bik, debtor.bank_name, debtor.reg_organization, debtor.reg_date, debtor.make_address_for_template_ul_ur(), debtor.make_address_for_template_ul_post())
                tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
                tables[2].rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                document.save('D:/python/FinMon/{}.docx'.format(last_operation.make_filename(date)))


def get_all_contracts_depending_on_sides(side1, side2, side3):
    print(side1, side2, side3)
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    exec_string = """SELECT contract_number
                        FROM contracts
                        JOIN organizations as side1 ON side1.org_id = contracts.side_1_id
                        JOIN organizations as side2 ON side2.org_id = contracts.side_2_id
                        JOIN organizations as side3 ON side3.org_id = contracts.side_3_id
                        WHERE (side1.org_name = :side1 OR side1.org_name = :side2 OR side1.org_name = :side3) 
                        AND (side2.org_name = :side1 OR side2.org_name = :side2 OR side2.org_name = :side3)
                        AND (side3.org_name = :side1 OR side3.org_name = :side2 OR side3.org_name = :side3)"""
    try:
        cur.execute(exec_string, {'side1': side1, 'side2': side2, 'side3': side3})
    except:
        print("JOPA", type(side1), side2, side3)
    rows = cur.fetchall()
    conn.close()
    contracts = []
    for row in rows:
        contracts.append(row[0])
    return contracts


def get_first_contract_depending_on_sides_ID(side1, side2, side3):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    exec_string = """SELECT contract_number
                        FROM contracts
                        JOIN organizations as side1 ON side1.org_id = contracts.side_1_id
                        JOIN organizations as side2 ON side2.org_id = contracts.side_2_id
                        JOIN organizations as side3 ON side3.org_id = contracts.side_3_id
                        WHERE (side1.org_id = :side1 OR side1.org_id = :side2 OR side1.org_id = :side3) 
                        AND (side2.org_id = :side1 OR side2.org_id = :side2 OR side2.org_id = :side3)
                        AND (side3.org_id = :side1 OR side3.org_id = :side2 OR side3.org_id = :side3)"""
    try:
        cur.execute(exec_string, {'side1': side1, 'side2': side2, 'side3': side3})
    except:
        print("JOPA", type(side1), side2, side3)
    rows = cur.fetchall()
    conn.close()
    contract = ""
    for row in rows:
        contract = row[0]
    return contract



def get_all_rs_depending_on_org_id(org_id):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT raschetniy_schet FROM accounts WHERE accounts.org_id = ?""", (org_id,))
    rows = cur.fetchall()
    conn.close()
    rs = []
    for row in rows:
        rs.append(row[0])
    return rs



def get_all_participants():
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT org_name FROM organizations""")
    rows = cur.fetchall()
    conn.close()
    orgs = []
    for row in rows:
        orgs.append(row[0])
    return orgs

def get_ur_participants():
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT org_name FROM organizations WHERE organizations.type_id = 1""")
    rows = cur.fetchall()
    conn.close()
    orgs = []
    for row in rows:
        orgs.append(row[0])
    return orgs

def get_fl_participants():
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT org_name FROM organizations WHERE organizations.type_id = 2""")
    rows = cur.fetchall()
    conn.close()
    orgs = []
    for row in rows:
        orgs.append(row[0])
    return orgs


def get_org_id_by_inn(inn):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    execute_string_ul = """SELECT org_id
                        from requisites_ul
                        WHERE requisites_ul.inn = ?"""
                       
    execute_string_fl = """SELECT org_id
                        from requisites_fl
                        WHERE requisites_fl.inn = ?"""
                 
    execute_string_ip = """SELECT org_id
                        from requisites_ip
                        WHERE requisites_ip.inn = ?"""

    fl = cur.execute(execute_string_fl, (inn,))
    fl_fetch = fl.fetchall()
    ul = cur.execute(execute_string_ul, (inn,))
    ul_fetch = ul.fetchall()
    ip = cur.execute(execute_string_ip, (inn,)) 
    ip_fetch = ip.fetchall()

    rows = fl_fetch
    if fl_fetch == []:
        rows = ul_fetch
        if ul_fetch == []:
            rows = ip_fetch       
    conn.close()
    org_id = ""
    for row in rows:
        org_id = row[0]
    return org_id


def get_org_id_by_short_name(short_name):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT org_id
                    from organizations
                    WHERE organizations.short_name =  ?""",(short_name,))
    rows = cur.fetchall()
    conn.close()
    org_id = ""
    for row in rows:
        org_id = row[0]
    return org_id


def get_ip_participants():
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT org_name FROM organizations WHERE organizations.type_id = 0""")
    rows = cur.fetchall()
    conn.close()
    orgs = []
    for row in rows:
        orgs.append(row[0])
    return orgs

def get_all_payment_purposes(contract):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    try:
        cur.execute("""SELECT payment_purpose FROM contracts WHERE contracts.contract_number = :contract""", {'contract': str(contract)})
    except:
        print(contract, type(contract))
    rows = cur.fetchall()
    conn.close()
    payment_purposes = []
    for row in rows:
        payment_purposes.append(row[0])
    return payment_purposes
        



        
def insert_operation(operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class,
                     payment_number, payer, debtor, receiver, contract):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""INSERT INTO operations (operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class,
                     payment_number, payer, debtor, receiver, contract) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""", (str(operation_date), str(detection_date), int(daily_number), int(annual_number), str(operation_number), float(amount), str(payment_purpose), str(payment_class),
                     str(payment_number), int(payer), int(debtor), int(receiver), str(contract)),)
    conn.commit()
    cur.close()
    conn.close()


def insert_new_organization(org_name, type_id):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""INSERT INTO organizations (org_name, type_id) VALUES (?,?)""", (str(org_name), str(type_id)),)
    conn.commit()
    cur.close()
    conn.close()

def insert_new_bank(bik, bank_name):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""INSERT INTO banks (bik, bank_name) VALUES (?,?)""", (str(bik), str(bank_name)),)
    conn.commit()
    cur.close()
    conn.close()

def get_org_id_by_org_name(org_name):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT org_id FROM organizations WHERE organizations.org_name = ?""", (org_name,))
    rows = cur.fetchall()
    conn.close()
    org_id = ''
    for row in rows:
        org_id = row[0]
    return org_id

def get_bik_by_bank_name(bank_name):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT bik FROM banks WHERE banks.bank_name = ?""", (bank_name,))
    rows = cur.fetchall()
    conn.close()
    bik = ''
    for row in rows:
        bik = row[0]
    return bik

def get_country_names():
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT country_name FROM country_codes_id""")
    rows = cur.fetchall()
    conn.close()
    country_names = []
    for row in rows:
        country_names.append(row[0])
    return sorted(country_names)

def get_okato():
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT okato_name FROM okato_subjects_id""")
    rows = cur.fetchall()
    conn.close()
    country_names = []
    for row in rows:
        country_names.append(row[0])
    return sorted(country_names)

def get_country_id_by_country_name(country_name):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT country_code_id FROM country_codes_id WHERE country_codes_id.country_name = ?""", (country_name,))
    rows = cur.fetchall()
    conn.close()
    country_name = ''
    for row in rows:
        country_name = row[0]
    return country_name


def get_okato_subject_id_by_okato_name(okato_name):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT okato_subject_id FROM okato_subjects_id WHERE okato_subjects_id.okato_name = ?""", (okato_name,))
    rows = cur.fetchall()
    conn.close()
    okato_name = ''
    for row in rows:
        okato_name = row[0]
    return okato_name


def insert_new_account(raschetniy_schet, bik, org_id):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""INSERT INTO accounts (raschetniy_schet, bik, org_id) VALUES (?,?,?)""", (str(raschetniy_schet), str(bik), str(org_id),))
    conn.commit()
    cur.close()
    conn.close()
    
def insert_new_contract(contract_number, contract_date, contract_description, payment_purpose, payment_class, side_1_id, side_2_id, side_3_id):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""INSERT INTO contracts (contract_number, contract_date, contract_description, payment_purpose,
                                         payment_class, side_1_id, side_2_id, side_3_id) VALUES (?,?,?,?,?,?,?,?)""", (str(contract_number), str(contract_date), str(contract_description), str(payment_purpose), str(payment_class), str(side_1_id), str(side_2_id), str(side_3_id),))
    conn.commit()
    cur.close()
    conn.close()    
    
    
    
    
def insert_new_ul_requisites(org_id, inn, kpp, okpo, okved, ogrn, reg_organization, reg_date, country_id_ur, okato_subject_id_ur,
                             area_ur, locality_ur, street_ur, house_ur, corpus_ur, office_ur, country_id_post, okato_subject_id_post, area_post,
                             locality_post, street_post, house_post, corpus_post, office_post):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""INSERT INTO requisites_ul (org_id, inn, kpp, okpo, okved, ogrn, reg_organization, reg_date, country_id_ur, okato_subject_id_ur,
                             area_ur, locality_ur, street_ur, house_ur, corpus_ur, office_ur, country_id_post, okato_subject_id_post, area_post,
                             locality_post, street_post, house_post, corpus_post, office_post) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", (org_id, inn, kpp, okpo, okved, ogrn, reg_organization, reg_date, country_id_ur, okato_subject_id_ur,
                             area_ur, locality_ur, street_ur, house_ur, corpus_ur, office_ur, country_id_post, okato_subject_id_post, area_post,
                             locality_post, street_post, house_post, corpus_post, office_post,))
    conn.commit()
    cur.close()
    conn.close()     


def insert_new_fl_requisites(org_id, first_name, middle_name, last_name, inn, country_citizenship_id, birth_date, birth_place_country_id, birth_place_okato_subject_id,
                             document_type, document_series, document_number, document_date, document_given_by, 
                             department_number, country_id_registration, okato_subject_id_registration, area_registration, locality_registration, 
                             street_registration, house_registration, corpus_registration, office_registration, country_id_residence, 
                             okato_subject_id_residence, area_residence, locality_residence, street_residence, house_residence, corpus_residence, office_residence):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""INSERT INTO requisites_fl (org_id, first_name, middle_name, last_name, inn, country_citizenship_id, birth_date, birth_place_country_id, birth_place_okato_subject_id,
                             document_type, document_series, document_number, document_date, document_given_by, 
                             department_number, country_id_registration, okato_subject_id_registration, area_registration, locality_registration, 
                             street_registration, house_registration, corpus_registration, office_registration, country_id_residence, 
                             okato_subject_id_residence, area_residence, locality_residence, street_residence, house_residence, corpus_residence, office_residence) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", (org_id, first_name, middle_name, last_name, inn, country_citizenship_id, birth_date, birth_place_country_id, birth_place_okato_subject_id,
                             document_type, document_series, document_number, document_date, document_given_by, department_number, country_id_registration, okato_subject_id_registration, area_registration, locality_registration, street_registration, house_registration, corpus_registration, office_registration, country_id_residence, 
                             okato_subject_id_residence, area_residence, locality_residence, street_residence, house_residence, corpus_residence, office_residence,))
    conn.commit()
    cur.close()
    conn.close()   

def insert_new_ip_requisites(org_id, first_name, middle_name, last_name, inn, ogrn, reg_organization, reg_date, country_citizenship_id, 
                             birth_date, birth_place_country_id, birth_place_okato_subject_id,document_type,
                             document_series, document_number, document_date, document_given_by, department_number, country_id_registration, okato_subject_id_registration,
                             area_registration, locality_registration, street_registration,house_registration, corpus_registration, office_registration,
                             country_id_residence, okato_subject_id_residence, area_residence, locality_residence, street_residence, house_residence, corpus_residence, office_residence):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""INSERT INTO requisites_ip (org_id, first_name, middle_name, last_name, inn, ogrn, reg_organization, reg_date, country_citizenship_id, 
                             birth_date, birth_place_country_id, birth_place_okato_subject_id,document_type,
                             document_series, document_number, document_given_by, department_number, document_date, country_id_registration, okato_subject_id_registration,
                             area_registration, locality_registration, street_registration,house_registration, corpus_registration, office_registration,
                             country_id_residence, okato_subject_id_residence, area_residence, locality_residence, street_residence, house_residence, corpus_residence, office_residence) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", (org_id, first_name, middle_name, last_name, inn, ogrn, reg_organization, reg_date, country_citizenship_id, 
                             birth_date, birth_place_country_id, birth_place_okato_subject_id,document_type,
                             document_series, document_number, document_date, document_given_by, department_number, country_id_registration, okato_subject_id_registration,
                             area_registration, locality_registration, street_registration,house_registration, corpus_registration, office_registration,
                             country_id_residence, okato_subject_id_residence, area_residence, locality_residence, street_residence, house_residence, corpus_residence, office_residence,))
    conn.commit()
    cur.close()
    conn.close()  

def get_type_id_by_type_name(type_name):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT type_id FROM org_types WHERE org_types.type_name = ?""", (type_name,))
    rows = cur.fetchall()
    conn.close()
    type_id = ''
    for row in rows:
        type_id = row[0]
    return type_id

def get_org_types():
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT type_name FROM org_types""")
    rows = cur.fetchall()
    conn.close()
    org_types = []
    for row in rows:
        org_types.append(row[0])
    return org_types

def get_banks():
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT bank_name FROM banks""")
    rows = cur.fetchall()
    conn.close()
    banks = []
    for row in rows:
        banks.append(row[0])
    return banks

def make_template_ul(start=None, end=None, num_operations=None):
    with open('D:/python/FinMon/template_ul.docx', 'rb') as file:
                document = Document(file)
                paragraphs = document.paragraphs
                tables = document.tables
                conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
                cur = conn.cursor()   
                query = ""
                
                if num_operations != None:
                    query = """SELECT * FROM (SELECT * FROM operations ORDER BY operations.operation_id DESC LIMIT ?) 
                                as r ORDER BY r.operation_id"""
                    cur.execute(query, (num_operations,))
                else:
                    query = """SELECT * FROM operations WHERE operation_id BETWEEN ? AND ?"""
                    cur.execute(query, (start, end,))
                    
                rows = cur.fetchall()
                conn.close()
                    
                for row in rows:
                    operation_object = OperationObject(*row[1:15])
                    payer_id = operation_object.payer
                    debtor_id = operation_object.debtor
                    receiver_id = operation_object.receiver
                        
                    payer = operation_object.get_requisites_by_type_id(participant=operation_object.payer)
                    debtor = operation_object.get_requisites_by_type_id(participant=operation_object.debtor)
                    receiver = operation_object.get_requisites_by_type_id(participant=operation_object.receiver)
                                            
                    paragraphs[1].text = 'ВНУТРЕННЕЕ СООБЩЕНИЕ (УВЕДОМЛЕНИЕ)\nо совершённой операции, подлежащей обязательному контролю, либо необычной (подозрительной) сделке № {} / {}'.format('{:08d}'.format(int(operation_object.daily_number)), '{:08d}'.format(int(operation_object.annual_number)))
                    paragraphs[1].runs[0].font.size = Pt(14)
                    paragraphs[1].runs[0].font.bold = True
                    paragraphs[1].runs[0].font.name = 'Times New Roman'
                        
                    paragraphs[2].text = """Мною, Максименко Егором Вячеславовичем, ведущим инженером ООО «РЦ ЮО», в ходе осуществления своей профессиональной деятельности {}. выявлена операция, совершенная {}, связанная с поступлением денежных средств на р/с{} {} от {} на основании платежного документа №{} на сумму {} (плательщик {}, основание платежа: {})""".format(operation_object.operation_date, operation_object.detection_date, receiver.account, receiver.org_name, payer.org_name, operation_object.payment_number, operation_object.amount, payer.org_name, operation_object.payment_purpose)
                    paragraphs[2].runs[0].font.name = 'Times New Roman'
                    paragraphs[2].runs[0].font.size = Pt(12)
                        
                    paragraphs[5].text = '{}'.format(receiver.org_name)
                    paragraphs[5].runs[0].font.size = Pt(12)
                    paragraphs[5].runs[0].font.bold = True
                    paragraphs[5].runs[0].font.name = 'Times New Roman'
                    paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    paragraphs[6].text = '{}'.format(payer.org_name)
                    paragraphs[6].runs[0].font.size = Pt(12)
                    paragraphs[6].runs[0].font.bold = True
                    paragraphs[6].runs[0].font.name = 'Times New Roman'
                    paragraphs[6].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    paragraphs[7].text = '{}'.format(debtor.org_name)
                    paragraphs[7].runs[0].font.size = Pt(12)
                    paragraphs[7].runs[0].font.bold = True
                    paragraphs[7].runs[0].font.name = 'Times New Roman'
                    paragraphs[7].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    tables[1].rows[0].cells[0].paragraphs[0].text = '{}г.'.format(operation_object.detection_date)
                    tables[1].rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                    tables[1].rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                        
                    tables[2].rows[0].cells[0].paragraphs[0].text = """{}\nИНН: {}\nКПП: {}\nОГРН: {}\nОКПО: {}\nОКВЭД: {}\nР/c: {}\nБИК: {}\nНаименование банка: {}\nНаименование регистрирующего органа: {}\nДата регистрации: {}\nЮридический адрес: {}\nФактический адрес: {}""".format(payer.org_name, payer.inn, payer.kpp, payer.ogrn, payer.okpo, payer.okved, payer.account, payer.bik, payer.bank_name, payer.reg_organization, payer.reg_date, payer.make_address_for_template_ul_ur(), payer.make_address_for_template_ul_post())
                    tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                    tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
                    tables[2].rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    tables[2].rows[0].cells[1].paragraphs[0].text = """{}\nИНН: {}\nКПП: {}\nОГРН: {}\nОКПО: {}\nОКВЭД: {}\nР/c: {}\nБИК: {}\nНаименование банка: {}\nНаименование регистрирующего органа: {}\nДата регистрации: {}\nЮридический адрес: {}\nФактический адрес: {}""".format(debtor.org_name, debtor.inn, debtor.kpp, debtor.ogrn, debtor.okpo, debtor.okved, debtor.account, debtor.bik, debtor.bank_name, debtor.reg_organization, debtor.reg_date, debtor.make_address_for_template_ul_ur(), debtor.make_address_for_template_ul_post())
                    tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                    tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
                    tables[2].rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    old_date = datetime.datetime.strptime(operation_object.detection_date, '%d.%m.%Y')
                    new_date = old_date.strftime('%Y%m%d')
                    document.save('D:/python/FinMon/{}.docx'.format(operation_object.make_filename(new_date)))


def get_operation_info_with_org_names(operation_number):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT 
                    operation_id,
                    operation_date,
                    detection_date,
                    daily_number,
                    annual_number,
                    operation_number,
                    amount,
                    operations.payment_purpose,
                    operations.payment_class,
                    payment_number,
                    payer.org_name AS payer,
                    debtor.org_name AS debtor,
                    receiver.org_name AS receiver,
                    contract,
                    contracts.contract_date,
                    contracts.contract_description,
                    debtor as debtor_id,
                    payer as payer_id,
                    receiver as receiver_id,
                    debtor.participant_type_id as debtor_participant,
					payer.participant_type_id as payer_participant,
					receiver.participant_type_id as receiver_participant
                    FROM operations
                    JOIN organizations as debtor ON debtor.org_id = operations.debtor
                    JOIN organizations as payer ON payer.org_id = operations.payer
                    JOIN organizations as receiver ON receiver.org_id = operations.receiver
                    JOIN contracts ON operations.contract = contracts.contract_number
                    WHERE operations.operation_number = ?
                    """,(operation_number,))
    rows = cur.fetchall()
    conn.close()
    operations = []
    for row in rows:
        operations.append(row)
    return rows[0]


def get_participant_id_by_org_id(org_id):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT participant_type_id
                    FROM 
                    organizations
                    WHERE organizations.org_id = ?
                    """,(org_id,))
    rows = cur.fetchall()
    conn.close()
    operations = []
    for row in rows:
        operations.append(row)
    return row[0]