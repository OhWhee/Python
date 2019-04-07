import xml.etree.ElementTree as ET
import re
from docx import Document

def read_xml(path):
    """Возваращает объект типа xml.etree.ElementTree"""
    tree = ET.parse(path)
    root = tree.getroot()
    return root
    
def get_operation_info(xml_root):
    """Создает класс с параметрами операции"""
    record_number = xml_root.findall(".//НомерЗаписи")[0].text
    operation_date = xml_root.findall(".//ДатаОперации")[0].text.replace('/', '.')
    detection_date = xml_root.findall(".//ДатаВыявления")[0].text.replace('/', '.')
    sum_operation = xml_root.findall(".//СумОперации")[0].text
    purpose = xml_root.findall(".//НазнПлатеж")[0].text
    pay_number = xml_root.findall(".//Коммент")[0].text
    operation_info_object = OperationInfo(record_number, operation_date,
                                             detection_date, sum_operation,
                                             purpose, pay_number)
    return operation_info_object


class OperationInfo:
    
    def __init__(self, record_number, operation_date, detection_date,
                 sum_operation, purpose, pay_number):
        self.record_number = record_number
        self.operation_date = operation_date
        self.detection_date = detection_date
        self.sum_operation = sum_operation
        self.purpose = purpose
        self.pay_number = pay_number
        
    def gatherAttrs(self):
        attrs = []
        for key in sorted(self.__dict__):
            attrs.append('%s: %s' % (key, getattr(self, key)))
        return ', \n'.join(attrs)

    def __str__(self):
        return '[%s: %s]' % (self.__class__.__name__, self.gatherAttrs())

class ParticipantUl():
    
    def __init__(self, org_name, org_inn, org_kpp, org_ogrn, org_okpo,
                 org_okved, org_rs, bik, bank_name, register_name, 
                 register_date, ur_adress, post_adress):
        self.org_name = org_name
        self.org_inn = org_inn
        self.org_kpp = org_kpp
        self.org_ogrn = org_ogrn
        self.org_okpo = org_okpo
        self.org_okved = org_okved
        self.org_rs = org_rs
        self.bik = bik
        self.bank_name = bank_name
        self.register_name = register_name
        self.register_date = register_date
        self.ur_adress = ur_adress
        self.post_adress = post_adress
        
    def gatherAttrs(self):
        attrs = []
        for key in sorted(self.__dict__):
            attrs.append('%s: %s' % (key, getattr(self, key)))
        return ', \n'.join(attrs)

    def __str__(self):
        return '[%s: %s]' % (self.__class__.__name__, self.gatherAttrs())

class ParticipantFl():
    
    def __init__(self, second_name, first_name, patronimic, inn, birth_date,
                 adress_reg, adress_preb, bik, bank_name, rs, passport_data):
        self.second_name = second_name
        self.first_name = first_name
        self.patronimic = patronimic
        self.inn = inn
        self.birth_date = birth_date
        self.adress_reg = adress_reg
        self.adress_preb = adress_preb
        self.bik = bik
        self.bank_name = bank_name
        self.rs = rs
        self.passport_data = passport_data
        
    def gatherAttrs(self):
        attrs = []
        for key in sorted(self.__dict__):
            attrs.append('%s: %s' % (key, getattr(self, key)))
        return ', \n'.join(attrs)

    def __str__(self):
        return '[%s: %s]' % (self.__class__.__name__, self.gatherAttrs())




def get_requisites(root, role):
    """Роли: Recipient - получатель
             Debtor - должник
             Payer - плательщик"""
    role_dict = {"Payer": "01", "Recipient": "02", "Debtor": "08"}
    role_code = role_dict.get(role)
    data = []
    for tags in root.findall(".//УчастникОп"):
        for tag in tags.getchildren():
            if tag.tag == "КодРоли" and tag.text == role_code:
                for tag in tags.getchildren():
                    if tag.tag == "СведЮЛ":
                        for i in tag:
                            #if i.text != re.match("\\n+\s+", str(i.text)) is None:
                            #print(tag.tag, i.tag, i.text)
                            #if i.text == None:
                            if i.text != re.match("\\n+\s+", str(i.text)):
                                i.text == ""
                            elif i.text == None:
                                i.text = ""
                            data.append({tags.tag + i.tag: i.text})
                    if tag.tag == "СведенияКО":
                        for k in tag:
                            #if k.text != re.match("\\n+\s+", str(k.text)) is None:
                            #print(tag.tag, k.tag, k.text)
                            #if k.text == None:
                            if k.text == re.match("\\n+\s+", str(k.text)):
                                k.text = ""
                            elif k.text == None:
                                k.text = ""
                            data.append({tags.tag + k.tag: k.text})
                    for t in tag.getchildren():
                        if t.tag == "ЮрАдр":
                            for x in t.getchildren():
                                #print(t.tag, x.tag)
                                if x.text == None:
                                    x.text = ""
                                data.append({t.tag + x.tag: x.text})
                        if t.tag == "ФактАдр":
                            for x in t.getchildren():
                                #print(t.tag, x.tag)
                                if x.text == None:
                                    x.text = ""
                                data.append({t.tag + x.tag: x.text})
    obj = combine_reqs_to_object(data)
    return obj

def combine_reqs_to_object(data):
    req_list = ['УчастникОпНаимЮЛ', 'УчастникОпИННЮЛ', 'УчастникОпКППЮЛ', 
                'УчастникОпОКПОЮЛ', 'УчастникОпОКВЭДЮЛ', 'УчастникОпОГРНЮЛ',
                'УчастникОпНаименРегОргана', 'УчастникОпДатаРегЮл', 'ЮрАдрПункт',
                'ЮрАдрУлица', 'ЮрАдрДом', 'ЮрАдрКорп', 'ЮрАдрОф', 'ФактАдрПункт',
                'ФактАдрУлица', 'ФактАдрДом', 'ФактАдрКорп', 'ФактАдрОф',
                'УчастникОпБИККО', 'УчастникОпНомерСчета', 'УчастникОпНаимКО']
    requisites = []
    req_list_index = 0
    for req in data:
        for key in req.keys():
            if key == req_list[req_list_index]:
                print(req_list_index)
                requisites.append(req.get(req_list[req_list_index]))
                if (req_list_index + 1) < (len(req_list)):
                    req_list_index = req_list_index + 1
    print(requisites)
    org_name = requisites[0]
    org_inn = requisites[1]
    org_kpp = requisites[2]
    org_ogrn = requisites[5]
    org_okpo = requisites[3]
    org_okved = requisites[4]
    org_rs = requisites[19]
    bik = requisites[18]
    bank_name = requisites[20]
    register_name = requisites[6]
    register_date = requisites[7].replace('/', '.')
    ur_adress = re.sub(" +", " ",'{} {} {} {} {}'.format(requisites[8], requisites[9], requisites[10],
                     requisites[11], requisites[12]))
    post_adress = re.sub(" +", " ", '{} {} {} {} {}'.format(requisites[13], requisites[14], requisites[15],
                     requisites[16], requisites[17]))
    participant = ParticipantUl(org_name, org_inn, org_kpp, org_ogrn, org_okpo,
                     org_okved, org_rs, bik, bank_name, register_name, 
                     register_date, ur_adress, post_adress)
    return participant






# TESTING
if __name__ == '__main__':
    test = read_xml('D:/RFM/xml monitoring(2018)/FM01_2_3444213503344401001_20181218_0100000003.xml')
    operation_info = get_operation_info(test)
    recipient_reqs = get_requisites(test, role="Recipient")
    debtor_reqs = get_requisites(test, role="Debtor")
    payer_reqs = get_requisites(test, role="Payer")


f = open('D:/python/FinMon/template.docx', 'rb')
document = Document(f)
tables = document.tables


def insertDataIntoParagraphs(document):
    for paragraph in document.paragraphs:
        if paragraph.text == 'fn.TransactionDetection':
            paragraph.text = operation_info.detection_date
        if paragraph.text == 'fn.TransactionDateг':
            paragraph.text = operation_info.detection_date
        if paragraph.text == 'fn.rsReceiver':
            paragraph.text = recipient_reqs.org_rs
        if paragraph.text == 'fn.nameReceiver':
            paragraph.text = recipient_reqs.org_name
        if paragraph.text =='fn.namePayer':
            paragraph.text = payer_reqs.org_name
        if paragraph.text =='fn.PayNumber':
            paragraph.text = operation_info.pay_number
        if paragraph.text =='fn.Payment':
            paragraph.text = operation_info.sum_operation
        if paragraph.text =='fn.namePayer':
            paragraph.text = payer_reqs.org_name
        if paragraph.text =='fn.Text':
            paragraph.text = operation_info.purpose
            
def insertDataIntoTables(tables):
    tables[2].rows[0].cells[0].text = payer_reqs.org_name
    tables[2].rows[1].cells[0].text = payer_reqs.org_inn
    tables[2].rows[2].cells[0].text = payer_reqs.org_kpp
    tables[2].rows[3].cells[0].text = payer_reqs.org_ogrn
    tables[2].rows[4].cells[0].text = payer_reqs.org_okpo
    tables[2].rows[5].cells[0].text = payer_reqs.org_okved
    tables[2].rows[6].cells[0].text = payer_reqs.org_rs
    tables[2].rows[7].cells[0].text = payer_reqs.org_bik
    tables[2].rows[8].cells[0].text = payer_reqs.bank_name
    tables[2].rows[9].cells[0].text = payer_reqs.register_name
    tables[2].rows[10].cells[0].text = payer_reqs.register_date 
    tables[2].rows[11].cells[0].text = payer_reqs.ur_adress
    tables[2].rows[12].cells[0].text = payer_reqs.post_adress
    
    tables[2].rows[0].cells[1].text = debtor_reqs.org_name
    tables[2].rows[1].cells[1].text = debtor_reqs.org_inn
    tables[2].rows[2].cells[1].text = debtor_reqs.org_kpp
    tables[2].rows[3].cells[1].text = debtor_reqs.org_ogrn
    tables[2].rows[4].cells[1].text = debtor_reqs.org_okpo
    tables[2].rows[5].cells[1].text = debtor_reqs.org_okved
    tables[2].rows[6].cells[1].text = debtor_reqs.org_rs
    tables[2].rows[7].cells[1].text = debtor_reqs.org_bik
    tables[2].rows[8].cells[1].text = debtor_reqs.bank_name
    tables[2].rows[9].cells[1].text = debtor_reqs.register_name
    tables[2].rows[10].cells[1].text = debtor_reqs.register_date 
    tables[2].rows[11].cells[1].text = debtor_reqs.ur_adress
    tables[2].rows[12].cells[1].text = debtor_reqs.post_adress

f.close()
