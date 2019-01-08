from docx import Document

f = open('D:/python/FinMon/template.docx', 'rb')
document = Document(f)
tables = document.tables
f.close()

def insertDataIntoParagraphs(document):
    for paragraph in document.paragraphs:
        if paragraph.text == 'fn.TransactionDetection':
            paragraph.text = operation_info.detection_date
        if paragraph.text == 'fn.TransactionDateг':
            paragraph.text = operation_info.detection_date
        if paragraph.text == 'fn.rsReceiver':
            paragraph.text = recipient_reqs.org_rs
        if paragraph.text == 'fn.nameReceiver':
            paragraph.text = recipient_reqs.org_name
        if paragraph.text =='fn.namePayer':
            paragraph.text = payer_reqs.org_name
        if paragraph.text =='fn.PayNumber':
            paragraph.text = operation_info.pay_number
        if paragraph.text =='fn.Payment':
            paragraph.text = operation_info.sum_operation
        if paragraph.text =='fn.namePayer':
            paragraph.text = payer_reqs.org_name
        if paragraph.text =='fn.Text':
            paragraph.text = operation_info.purpose
            
def insertDataIntoTables(tables):
    tables[2].rows[0].cells[0].text = payer_reqs.org_name
    tables[2].rows[1].cells[0].text = payer_reqs.org_inn
    tables[2].rows[2].cells[0].text = payer_reqs.org_kpp
    tables[2].rows[3].cells[0].text = payer_reqs.org_ogrn
    tables[2].rows[4].cells[0].text = payer_reqs.org_okpo
    tables[2].rows[5].cells[0].text = payer_reqs.org_okved
    tables[2].rows[6].cells[0].text = payer_reqs.org_rs
    tables[2].rows[7].cells[0].text = payer_reqs.org_bik
    tables[2].rows[8].cells[0].text = payer_reqs.bank_name
    tables[2].rows[9].cells[0].text = payer_reqs.register_name
    tables[2].rows[10].cells[0].text = payer_reqs.register_date 
    tables[2].rows[11].cells[0].text = payer_reqs.ur_adress
    tables[2].rows[12].cells[0].text = payer_reqs.post_adress
    
    tables[2].rows[0].cells[1].text = debtor_reqs.org_name
    tables[2].rows[1].cells[1].text = debtor_reqs.org_inn
    tables[2].rows[2].cells[1].text = debtor_reqs.org_kpp
    tables[2].rows[3].cells[1].text = debtor_reqs.org_ogrn
    tables[2].rows[4].cells[1].text = debtor_reqs.org_okpo
    tables[2].rows[5].cells[1].text = debtor_reqs.org_okved
    tables[2].rows[6].cells[1].text = debtor_reqs.org_rs
    tables[2].rows[7].cells[1].text = debtor_reqs.org_bik
    tables[2].rows[8].cells[1].text = debtor_reqs.bank_name
    tables[2].rows[9].cells[1].text = debtor_reqs.register_name
    tables[2].rows[10].cells[1].text = debtor_reqs.register_date 
    tables[2].rows[11].cells[1].text = debtor_reqs.ur_adress
    tables[2].rows[12].cells[1].text = debtor_reqs.post_adress

    
    