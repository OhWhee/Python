# -*- coding: utf-8 -*-
import sqlite3
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#debtor = ('8',)
#receiver = ('17',)
#payer = ('7',)
#
#participants = {'debtor' : '8', 'receiver' : '17', 'payer' : '7'}

class ParticipantObject():
    
    def __init__(self, org_name, type_id, inn, kpp, okpo, okved, ogrn, reg_organization, reg_date, country_id_ur, okato_subject_id_ur, area_ur, locality_ur,
                 street_ur, house_ur, corpus_ur, office_ur, country_id_post, okato_subject_id_post, area_post, locality_post, street_post, 
                 house_post, corpus_post, office_post, raschetniy_schet, bik, bank_name):
        self.org_name = org_name
        self.type_id = type_id
        self.inn = inn
        self.kpp = kpp
        self.okpo = okpo
        self.okved = okved
        self.ogrn = ogrn
        self.reg_organization = reg_organization
        self.reg_date = reg_date
        self.country_id_ur = country_id_ur
        self.okato_subject_id_ur = okato_subject_id_ur
        self.area_ur = area_ur
        self.locality_ur = locality_ur
        self.street_ur = street_ur
        self.house_ur = house_ur
        self.corpus_ur = corpus_ur
        self.office_ur = office_ur
        self.country_id_post = country_id_post
        self.okato_subject_id_post = okato_subject_id_post
        self.area_post = area_post
        self.locality_post = locality_post
        self.street_post = street_post
        self.house_post = house_post
        self.corpus_post = corpus_post
        self.office_post = office_post
        self.raschetniy_schet = raschetniy_schet
        self.bik = bik
        self.bank_name = bank_name
        
    def make_address_for_template_ul_ur(self):
        return '{} {} {} {} {}'.format(self.locality_ur, self.street_ur, self.house_ur, self.corpus_ur, self.office_ur)
    
    def make_address_for_template_ul_post(self):
        return '{} {} {} {} {}'.format(self.locality_post, self.street_post, self.house_post, self.corpus_post, self.office_post)


class OperationObjectFactory():
        
        def create_last_operation_object(self):
            conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
            cur = conn.cursor()
            annual_number = ''
            cur.execute("""SELECT * FROM operations WHERE ROWID IN (SELECT max(ROWID) FROM operations)""")
            rows = cur.fetchall()
            operation_date = ''
            detection_date = ''
            daily_number = ''
            annual_number = ''
            operation_number = ''
            amount = ''
            payment_purpose = ''
            payment_class = ''
            payment_number = ''
            debtor = ''
            payer = ''
            receiver = ''
            contract = ''
            operation_list = [operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract]
            i=0
            for r in rows[0][1:]:
                operation_list[0+i] = r
                i+=1
            operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract = [str(e) for e in operation_list]
            return OperationObject(operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract)  

        def create_operation_object_by_id(self, operation_id):
            conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
            cur = conn.cursor()
            annual_number = ''
            cur.execute("""SELECT * FROM operations WHERE operation_id = ?""", (operation_id,))
            rows = cur.fetchall()
            operation_date = ''
            detection_date = ''
            daily_number = ''
            annual_number = ''
            operation_number = ''
            amount = ''
            payment_purpose = ''
            payment_class = ''
            payment_number = ''
            debtor = ''
            payer = ''
            receiver = ''
            contract = ''
            operation_list = [operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract]
            i=0
            for r in rows[0][1:]:
                operation_list[0+i] = r
                i+=1
            operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract = [str(e) for e in operation_list]
            return OperationObject(operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class, payment_number, debtor, payer, receiver, contract)


        


class OperationObject():
    
    def __init__(self, operation_date, detection_date, daily_number, annual_number,
                 operation_number, amount, payment_purpose, payment_class, payment_number, payer,
                 debtor, receiver, contract):
        self.operation_date = operation_date
        self.detection_date = detection_date
        self.daily_number = daily_number
        self.annual_number = annual_number
        self.operation_number = operation_number
        self.amount = amount
        self.payment_purpose = payment_purpose
        self.payment_class = payment_class
        self.payment_number = payment_number
        self.payer = payer
        self.debtor = debtor
        self.receiver = receiver
        self.contract = contract
        
        
    def get_requisites_for_ul(self, participant):
        conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
        cur = conn.cursor()
        cur.execute("""SELECT 
                            organizations.org_name,
                            organizations.type_id,
                            requisites_ul.inn,
                            requisites_ul.kpp,
                            requisites_ul.okpo,
                            requisites_ul.okved,
                            requisites_ul.ogrn,
                            requisites_ul.reg_organization,
                            requisites_ul.reg_date,
                            requisites_ul.country_id_ur,
                            requisites_ul.okato_subject_id_ur,
                            requisites_ul.area_ur,
                            requisites_ul.locality_ur,
                            requisites_ul.street_ur,
                            requisites_ul.house_ur,
                            requisites_ul.corpus_ur,
                            requisites_ul.office_ur,
                            requisites_ul.country_id_post,
                            requisites_ul.okato_subject_id_post,
                            requisites_ul.area_post,
                            requisites_ul.locality_post,
                            requisites_ul.street_post,
                            requisites_ul.house_post,
                            requisites_ul.corpus_post,
                            requisites_ul.office_post,
                            accounts.raschetniy_schet,
                            accounts.bik,
                            banks.bank_name
                            FROM organizations
                            JOIN requisites_ul ON organizations.org_id = requisites_ul.org_id
                            JOIN accounts ON organizations.org_id = accounts.org_id
                            JOIN banks ON banks.bik = accounts.bik
                            WHERE organizations.org_id = ?""", (participant,))
        rows = cur.fetchall()
        obj = ParticipantObject(*rows[0])
        return obj      
        
        
    def make_next_annual_number(self):
        print('Предыдущий номер ФЭС - {}, следующий - {}'.format(self.annual_number, str(int(self.annual_number)+1)))
        return str(int(self.annual_number)+1)


    def make_next_operation_number(self):
        annual_number = int(self.make_next_annual_number())
        now = datetime.datetime.now()
        formatter = '{:08d}'.format
        operation_number = '{}_3444213503_344401001_{}'.format(now.year, formatter(annual_number))
        return operation_number

    def make_filename(self, date):
        daily_number = int(self.daily_number)
#        now = datetime.datetime.now()
        formatter = '{:08d}'.format
        filename = 'FM01_2_3444213503344401001_{}_01{}'.format(date, formatter(daily_number))
        return filename

    def get_last_detection_date(self):
        return self.detection_date
    
    
    def make_new_daily_number(self):
        now = datetime.datetime.now()
        detection_date = self.get_last_detection_date()
        daily_number = int(self.get_last_daily_number())
        if str(now.strftime("%d.%m.%Y")) != detection_date:
            daily_number = 1
        else:
            daily_number += 1
        return daily_number


    def get_last_daily_number(self):
        return self.daily_number
    
    
    def make_template_ul(self, operation_id, date=datetime.datetime.now().strftime("%Y%m%d")):
            with open('D:/python/FinMon/template_ul.docx', 'rb') as file:
                document = Document(file)
                paragraphs = document.paragraphs
                tables = document.tables
                last_operation = self.create_operation_object_by_id(operation_id)
                payer_id = last_operation.payer
                debtor_id = last_operation.debtor
                receiver_id = last_operation.receiver
                
                payer = last_operation.get_requisites_for_ul(participant=payer_id)
                debtor = last_operation.get_requisites_for_ul(participant=debtor_id)
                receiver = last_operation.get_requisites_for_ul(participant=receiver_id)
                
                paragraphs[1].text = 'ВНУТРЕННЕЕ СООБЩЕНИЕ (УВЕДОМЛЕНИЕ)\nо совершённой операции, подлежащей обязательному контролю, либо необычной (подозрительной) сделке № {} / {}'.format('{:08d}'.format(int(last_operation.daily_number)), '{:08d}'.format(int(last_operation.annual_number)))
                paragraphs[1].runs[0].font.size = Pt(14)
                paragraphs[1].runs[0].font.bold = True
                paragraphs[1].runs[0].font.name = 'Times New Roman'
                
                paragraphs[2].text = """Мною, Максименко Егором Вячеславовичем, ведущим инженером ООО «РЦ ЮО», в ходе осуществления своей профессиональной деятельности {}. выявлена операция, совершенная {}, связанная с поступлением денежных средств на р/с{} {} от {} на основании платежного документа №{} на сумму {} (плательщик {}, основание платежа: {})""".format(last_operation.detection_date, last_operation.operation_date, receiver.raschetniy_schet, receiver.org_name, payer.org_name, last_operation.payment_number, last_operation.amount, payer.org_name, last_operation.payment_purpose)
                paragraphs[2].runs[0].font.name = 'Times New Roman'
                paragraphs[2].runs[0].font.size = Pt(12)
                
                paragraphs[5].text = '{}'.format(receiver.org_name)
                paragraphs[5].runs[0].font.size = Pt(12)
                paragraphs[5].runs[0].font.bold = True
                paragraphs[5].runs[0].font.name = 'Times New Roman'
                paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                paragraphs[6].text = '{}'.format(payer.org_name)
                paragraphs[6].runs[0].font.size = Pt(12)
                paragraphs[6].runs[0].font.bold = True
                paragraphs[6].runs[0].font.name = 'Times New Roman'
                paragraphs[6].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                paragraphs[7].text = '{}'.format(debtor.org_name)
                paragraphs[7].runs[0].font.size = Pt(12)
                paragraphs[7].runs[0].font.bold = True
                paragraphs[7].runs[0].font.name = 'Times New Roman'
                paragraphs[7].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                tables[1].rows[0].cells[0].paragraphs[0].text = '{}г.'.format(last_operation.detection_date)
                tables[1].rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                tables[1].rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                
                tables[2].rows[0].cells[0].paragraphs[0].text = """{}\nИНН: {}\nКПП: {}\nОГРН: {}\nОКПО: {}\nОКВЭД: {}\nР/c: {}\nБИК: {}\nНаименование банка: {}\nНаименование регистрирующего органа: {}\nДата регистрации: {}\nЮридический адрес: {}\nФактический адрес: {}""".format(payer.org_name, payer.inn, payer.kpp, payer.ogrn, payer.okpo, payer.okved, payer.raschetniy_schet, payer.bik, payer.bank_name, payer.reg_organization, payer.reg_date, payer.make_address_for_template_ul_ur(), payer.make_address_for_template_ul_post())
                tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
                tables[2].rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                tables[2].rows[0].cells[1].paragraphs[0].text = """{}\nИНН: {}\nКПП: {}\nОГРН: {}\nОКПО: {}\nОКВЭД: {}\nР/c: {}\nБИК: {}\nНаименование банка: {}\nНаименование регистрирующего органа: {}\nДата регистрации: {}\nЮридический адрес: {}\nФактический адрес: {}""".format(debtor.org_name, debtor.inn, debtor.kpp, debtor.ogrn, debtor.okpo, debtor.okved, debtor.raschetniy_schet, debtor.bik, debtor.bank_name, debtor.reg_organization, debtor.reg_date, debtor.make_address_for_template_ul_ur(), debtor.make_address_for_template_ul_post())
                tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
                tables[2].rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                document.save('D:/python/FinMon/{}.docx'.format(last_operation.make_filename(date)))


def get_all_contracts_depending_on_sides(side1, side2, side3):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    exec_string = """SELECT contract_number
                        FROM contracts
                        JOIN organizations as side1 ON side1.org_id = contracts.side_1_id
                        JOIN organizations as side2 ON side2.org_id = contracts.side_2_id
                        JOIN organizations as side3 ON side3.org_id = contracts.side_3_id
                        WHERE (side1.org_name = :side1 OR side1.org_name = :side2 OR side1.org_name = :side3) 
                        AND (side2.org_name = :side1 OR side2.org_name = :side2 OR side2.org_name = :side3)
                        AND (side3.org_name = :side1 OR side3.org_name = :side2 OR side3.org_name = :side3)"""
    try:
        cur.execute(exec_string, {'side1': side1, 'side2': side2, 'side3': side3})
    except:
        print("JOPA", type(side1), side2, side3)
    rows = cur.fetchall()
    conn.close()
    contracts = []
    for row in rows:
        contracts.append(row[0])
    return contracts


def get_all_participants():
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT org_name FROM organizations""")
    rows = cur.fetchall()
    conn.close()
    orgs = []
    for row in rows:
        orgs.append(row[0])
    return orgs

def get_all_payment_purposes(contract):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    try:
        cur.execute("""SELECT payment_purpose FROM contracts WHERE contracts.contract_number = :contract""", {'contract': str(contract)})
    except:
        print(contract, type(contract))
    rows = cur.fetchall()
    conn.close()
    payment_purposes = []
    for row in rows:
        payment_purposes.append(row[0])
    return payment_purposes
        
        
def insert_operation(operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class,
                     payment_number, payer, debtor, receiver, contract):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""INSERT INTO operations (operation_date, detection_date, daily_number, annual_number, operation_number, amount, payment_purpose, payment_class,
                     payment_number, payer, debtor, receiver, contract) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""", (str(operation_date), str(detection_date), int(daily_number), int(annual_number), str(operation_number), float(amount), str(payment_purpose), str(payment_class),
                     int(payment_number), int(payer), int(debtor), int(receiver), str(contract)),)
    conn.commit()
    cur.close()
    conn.close()


def get_org_id_by_org_name(org_name):
    conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
    cur = conn.cursor()
    cur.execute("""SELECT org_id FROM organizations WHERE organizations.org_name = ?""", (org_name,))
    rows = cur.fetchall()
    conn.close()
    org_id = ''
    for row in rows:
        org_id = row[0]
    return org_id


def make_template_ul(start=None, end=None, num_operations=None):
    with open('D:/python/FinMon/template_ul.docx', 'rb') as file:
                document = Document(file)
                paragraphs = document.paragraphs
                tables = document.tables
                conn = sqlite3.connect('D:\python\FinMon\FinMon.db')
                cur = conn.cursor()   
                query = ""
                
                if num_operations != None:
                    query = """SELECT * FROM (SELECT * FROM operations ORDER BY operations.operation_id DESC LIMIT ?) 
                                as r ORDER BY r.operation_id"""
                    cur.execute(query, (num_operations,))
                else:
                    query = """SELECT * FROM operations WHERE operation_id BETWEEN ? AND ?"""
                    cur.execute(query, (start, end,))
                    
                rows = cur.fetchall()
                conn.close()
                    
                for row in rows:
                    operation_object = OperationObject(*row[1:15])
                    payer_id = operation_object.payer
                    debtor_id = operation_object.debtor
                    receiver_id = operation_object.receiver
                        
                    payer = operation_object.get_requisites_for_ul(participant=payer_id)
                    debtor = operation_object.get_requisites_for_ul(participant=debtor_id)
                    receiver = operation_object.get_requisites_for_ul(participant=receiver_id)
                                            
                    paragraphs[1].text = 'ВНУТРЕННЕЕ СООБЩЕНИЕ (УВЕДОМЛЕНИЕ)\nо совершённой операции, подлежащей обязательному контролю, либо необычной (подозрительной) сделке № {} / {}'.format('{:08d}'.format(int(operation_object.daily_number)), '{:08d}'.format(int(operation_object.annual_number)))
                    paragraphs[1].runs[0].font.size = Pt(14)
                    paragraphs[1].runs[0].font.bold = True
                    paragraphs[1].runs[0].font.name = 'Times New Roman'
                        
                    paragraphs[2].text = """Мною, Максименко Егором Вячеславовичем, ведущим инженером ООО «РЦ ЮО», в ходе осуществления своей профессиональной деятельности {}. выявлена операция, совершенная {}, связанная с поступлением денежных средств на р/с{} {} от {} на основании платежного документа №{} на сумму {} (плательщик {}, основание платежа: {})""".format(operation_object.operation_date, operation_object.detection_date, receiver.raschetniy_schet, receiver.org_name, payer.org_name, operation_object.payment_number, operation_object.amount, payer.org_name, operation_object.payment_purpose)
                    paragraphs[2].runs[0].font.name = 'Times New Roman'
                    paragraphs[2].runs[0].font.size = Pt(12)
                        
                    paragraphs[5].text = '{}'.format(receiver.org_name)
                    paragraphs[5].runs[0].font.size = Pt(12)
                    paragraphs[5].runs[0].font.bold = True
                    paragraphs[5].runs[0].font.name = 'Times New Roman'
                    paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    paragraphs[6].text = '{}'.format(payer.org_name)
                    paragraphs[6].runs[0].font.size = Pt(12)
                    paragraphs[6].runs[0].font.bold = True
                    paragraphs[6].runs[0].font.name = 'Times New Roman'
                    paragraphs[6].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    paragraphs[7].text = '{}'.format(debtor.org_name)
                    paragraphs[7].runs[0].font.size = Pt(12)
                    paragraphs[7].runs[0].font.bold = True
                    paragraphs[7].runs[0].font.name = 'Times New Roman'
                    paragraphs[7].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    tables[1].rows[0].cells[0].paragraphs[0].text = '{}г.'.format(operation_object.detection_date)
                    tables[1].rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                    tables[1].rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                        
                    tables[2].rows[0].cells[0].paragraphs[0].text = """{}\nИНН: {}\nКПП: {}\nОГРН: {}\nОКПО: {}\nОКВЭД: {}\nР/c: {}\nБИК: {}\nНаименование банка: {}\nНаименование регистрирующего органа: {}\nДата регистрации: {}\nЮридический адрес: {}\nФактический адрес: {}""".format(payer.org_name, payer.inn, payer.kpp, payer.ogrn, payer.okpo, payer.okved, payer.raschetniy_schet, payer.bik, payer.bank_name, payer.reg_organization, payer.reg_date, payer.make_address_for_template_ul_ur(), payer.make_address_for_template_ul_post())
                    tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
                    tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    tables[2].rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
                    tables[2].rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    tables[2].rows[0].cells[1].paragraphs[0].text = """{}\nИНН: {}\nКПП: {}\nОГРН: {}\nОКПО: {}\nОКВЭД: {}\nР/c: {}\nБИК: {}\nНаименование банка: {}\nНаименование регистрирующего органа: {}\nДата регистрации: {}\nЮридический адрес: {}\nФактический адрес: {}""".format(debtor.org_name, debtor.inn, debtor.kpp, debtor.ogrn, debtor.okpo, debtor.okved, debtor.raschetniy_schet, debtor.bik, debtor.bank_name, debtor.reg_organization, debtor.reg_date, debtor.make_address_for_template_ul_ur(), debtor.make_address_for_template_ul_post())
                    tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
                    tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    tables[2].rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
                    tables[2].rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    old_date = datetime.datetime.strptime(operation_object.detection_date, '%d.%m.%Y')
                    new_date = old_date.strftime('%Y%m%d')
                    document.save('D:/python/FinMon/{}.docx'.format(operation_object.make_filename(new_date)))
